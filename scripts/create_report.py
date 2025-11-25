from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_paragraph(doc, text):
    p = doc.add_paragraph(text)
    p.style = doc.styles['Normal']
    return p

def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style='List Bullet')

def build_report():
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(11)

    title = doc.add_heading('Digital Item Catalog Viewer — Project Status Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, 'Executive Summary', level=1)
    add_paragraph(doc, 'We built a working Digital Item Catalog Viewer suitable for self‑checkout kiosks and store demos. The frontend delivers responsive, media‑rich browsing with search, categorization, pagination, and a simple cart/checkout flow. The backend exposes product data with server‑side pagination and a consolidated checkout endpoint. The system is verified with lint/build checks and is ready to demo to supermarkets and shops.')

    add_heading(doc, 'Objectives', level=1)
    add_bullets(doc, [
        'Present a visually engaging, responsive catalog for browsing products with images, descriptions, prices, and categories.',
        'Support kiosk workflows: quick add to cart, minimal friction checkout, idle attract overlay.',
        'Establish a backend API enabling scalable browsing (server‑side pagination) and consolidated checkout.'
    ])

    add_heading(doc, 'Deliverables', level=1)
    add_paragraph(doc, 'Frontend (React + Vite + TypeScript + Chakra UI v3)')
    add_bullets(doc, [
        'Search, category filters, and server‑side pagination for scalable catalogs.',
        'Product cards with images, category badge, and price; responsive, touch‑friendly UI.',
        'Item detail view with larger media and optional video.',
        'Cart and checkout posting a full cart payload to the backend; success banner with calculated total.',
        'Kiosk mode overlay: idle attract screen after inactivity, dismissible on touch/click.'
    ])
    add_paragraph(doc, 'Items Backend (FastAPI)')
    add_bullets(doc, [
        'Server‑side items pagination with query params for page, per_page, search, and category.',
        'Categories endpoint providing category list.',
        'Consolidated checkout endpoint accepting full cart payload and returning totals.',
        'CORS configured for local dev and static preview origins.'
    ])

    add_heading(doc, 'Key Features Implemented', level=1)
    add_bullets(doc, [
        'Interactive browsing: search and category chips, backend‑driven pagination.',
        'Multimedia: product images on list and detail views; optional video player.',
        'Responsive design: Chakra UI layouts optimized for kiosks and touch.',
        'Categorization: item categories and a categories API for filtering.',
        'Checkout flow: cart payload sent to backend; totals computed server‑side.'
    ])

    add_heading(doc, 'Architecture', level=1)
    add_bullets(doc, [
        'Frontend: React 19 + Vite 7, TypeScript 5.9, Chakra UI v3.',
        'Backend: FastAPI serving items, categories, item detail, submissions, and checkout.',
        'Optional services: Notify microservice scaffolded for future integrations.'
    ])

    add_heading(doc, 'Primary Endpoints', level=1)
    add_bullets(doc, [
        'GET /items?page=&per_page=&q=&category= → { items, total, page, per_page } (items-backend/main.py:46)',
        'GET /items/{id} → single item (items-backend/main.py:65)',
        'GET /categories → list of category names (items-backend/main.py:81)',
        'POST /checkout → accepts cart payload and returns total (items-backend/main.py:104)'
    ])

    add_heading(doc, 'Frontend Highlights', level=1)
    add_bullets(doc, [
        'Catalog page with category chips, images, prices, and server‑side pagination (catalog-frontend/src/pages/ItemsList.tsx:26).',
        'Product detail with larger media and optional video (catalog-frontend/src/pages/ItemDetail.tsx:58).',
        'Chakra provider and app shell (catalog-frontend/src/App.tsx:1).',
        'Kiosk attract overlay (catalog-frontend/src/components/KioskShell.tsx:1).',
        'API functions for items, categories, checkout (catalog-frontend/src/api.ts:1).'
    ])

    add_heading(doc, 'Quality Gates', level=1)
    add_bullets(doc, [
        'Linting: npm run lint passes.',
        'Build: npm run build succeeds and generates dist/.',
        'CORS updated to allow dev (5173) and preview (5174) origins.'
    ])

    add_heading(doc, 'How To Run', level=1)
    add_bullets(doc, [
        'Items backend: python -m uvicorn main:app --reload --port 8001 (c:/Work/ADG Intern/items-backend).',
        'Frontend dev: npm run dev → http://localhost:5173/ (c:/Work/ADG Intern/catalog-frontend).',
        'Frontend preview: npm run build && npx vite preview --port 5174 --host → http://localhost:5174/.'
    ])

    add_heading(doc, 'Demo Flow', level=1)
    add_bullets(doc, [
        'Open the catalog, use search and category chips to filter.',
        'Add items to cart, view total.',
        'Submit checkout; banner displays backend‑calculated total.'
    ])

    add_heading(doc, 'Decisions & Rationale', level=1)
    add_bullets(doc, [
        'Chakra UI v3 chosen for accessible and responsive kiosk UI.',
        'Server‑side pagination for performance and scalability.',
        'Consolidated checkout endpoint simplifies demo and integration.',
        'Idle attract overlay to match kiosk expectations and assist store‑floor demos.'
    ])

    add_heading(doc, 'Issues Resolved', level=1)
    add_bullets(doc, [
        'Chakra v3 prop/type changes: provider value, button disabled, HStack gap.',
        'React fast refresh rule by moving app tree into App.tsx.',
        'CORS expanded to allow both dev and preview origins.'
    ])

    add_heading(doc, 'Next Enhancements', level=1)
    add_bullets(doc, [
        'Department navigator (hierarchical categories with sidebar).',
        'Quantity selectors and per‑item modifiers on cards/detail.',
        'Image CDN optimization for faster thumbnails.',
        'Itemized receipts and order persistence (SQLite/Postgres).',
        'Kiosk features: full‑screen mode, attract slideshow, inactivity reset.'
    ])

    doc.save('Digital Item Catalog Viewer — Project Status Report.docx')

if __name__ == '__main__':
    build_report()

